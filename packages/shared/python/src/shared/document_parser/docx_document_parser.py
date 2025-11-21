"""Document parser implementation for multiple file formats."""

from abc import ABC, abstractmethod
from typing import BinaryIO, Dict, Any, List, Optional, Union

from shared.parsing.document_parser import DocumentParser
from shared.types.document import Document
from shared.types.common import PrimitiveValue, MetadataValue

try:
    import pypdf
    from docx import Document as DocxDocument
except ImportError:
    pypdf = None
    DocxDocument = None

import os
import tempfile
from pathlib import Path


class DocxDocumentParser(DocumentParser):
    """DOCX document parser using python-docx library."""

    def __init__(self):
        """Initialize DOCX parser."""
        pass

    def parse(self, raw: BinaryIO, metadata: dict[str, Any] | None = None) -> Document:
        """
        Parse DOCX document bytes into structured Document.

        Args:
            raw: Binary DOCX content
            metadata: Optional metadata to attach

        Returns:
            Document: Structured document with text and metadata

        Raises:
            DocumentParsingError: If DOCX parsing fails
        """
        if not DocxDocument:
            raise DocumentParsingError("python-docx library not available for DOCX parsing")

        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(raw.read())
                temp_file_path = temp_file.name
            
            try:
                # Open the document
                doc = DocxDocument(temp_file_path)
                
                # Extract text from paragraphs
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                text_content = "\n".join(text_parts)
                
                # Extract metadata from DOCX
                docx_metadata = self._extract_docx_metadata(doc, metadata)
                
                return Document(
                    id=self._generate_document_id(metadata),
                    workspace_id=metadata.get("workspace_id", "default") if metadata else "default",
                    title=metadata.get("title", self._extract_title_from_metadata(docx_metadata)) if metadata else self._extract_title_from_metadata(docx_metadata),
                    content=text_content,
                    metadata={
                        **docx_metadata,
                        "file_type": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "char_count": len(text_content),
                    }
                )
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}") from e

    def supports_format(self, filename: str) -> bool:
        """
        Check if parser supports the file format.

        Args:
            filename: Name of the file to check

        Returns:
            bool: True if DOCX format is supported
        """
        return filename.lower().endswith('.docx')

    def extract_metadata(self, raw: BinaryIO) -> dict[str, Any]:
        """
        Extract metadata from raw DOCX bytes.

        Args:
            raw: Binary DOCX content

        Returns:
            dict: Extracted metadata
        """
        if not DocxDocument:
            return {}
        
        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(raw.read())
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                return self._extract_docx_metadata(doc, None)
            finally:
                os.unlink(temp_file_path)
        except Exception:
            return {}

    def _extract_docx_metadata(self, doc: Any, user_metadata: dict[str, Any] | None) -> dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: python-docx document object
            user_metadata: User-provided metadata

        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata: dict[str, Any] = {}
        
        # Basic document info
        metadata.update({
            "paragraph_count": len(doc.paragraphs),
        })
        
        # Core properties
        if hasattr(doc, 'core_properties'):
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "category": core_props.category,
                    "comments": core_props.comments,
                    "created": core_props.created,
                    "modified": core_props.modified,
                    "last_modified_by": core_props.last_modified_by,
                    "revision": core_props.revision,
                })
        
        # Include user metadata
        if user_metadata:
            metadata.update(user_metadata)
        
        return metadata

    def _extract_title_from_metadata(self, metadata: dict[str, Any]) -> str | None:
        """Extract title from DOCX metadata."""
        # Try multiple fields for title
        title = metadata.get("title")
        if title:
            return str(title)
        
        title = metadata.get("subject")
        if title:
            return str(title)
        
        return None

    def _generate_document_id(self, metadata: dict[str, Any] | None) -> str:
        """Generate document ID from metadata or use default."""
        if metadata and "document_id" in metadata:
            return str(metadata["document_id"])
        
        import uuid
        return str(uuid.uuid4())
