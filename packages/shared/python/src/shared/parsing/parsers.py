"""Document parser implementation for multiple file formats."""

from abc import ABC, abstractmethod
from typing import BinaryIO, Dict, Any, List, Optional, Union

from shared.parsing.document_parser import DocumentParser
from shared.types.document import Document
from shared.types.common import PrimitiveValue, MetadataValue

try:
    import pypdf
    from docx import Document as DocxDocument
except ImportError:
    pypdf = None
    DocxDocument = None

import os
import tempfile
from pathlib import Path


class PDFDocumentParser(DocumentParser):
    """PDF document parser using PyPDF2 library."""

    def __init__(self):
        """Initialize PDF parser."""
        pass

    def parse(self, raw: BinaryIO, metadata: dict[str, Any] | None = None) -> Document:
        """
        Parse PDF document bytes into structured Document.

        Args:
            raw: Binary PDF content
            metadata: Optional metadata to attach

        Returns:
            Document: Structured document with text and metadata

        Raises:
            DocumentParsingError: If PDF parsing fails
        """
        if not pypdf:
            raise DocumentParsingError("PyPDF2 library not available for PDF parsing")

        try:
            reader = pypdf.PdfReader(raw)
            
            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            text_content = "\n".join(text_parts)
            
            # Extract metadata from PDF
            pdf_metadata = self._extract_pdf_metadata(reader, metadata)
            
            return Document(
                id=self._generate_document_id(metadata),
                workspace_id=metadata.get("workspace_id", "default") if metadata else "default",
                title=metadata.get("title", self._extract_title_from_metadata(pdf_metadata)) if metadata else self._extract_title_from_metadata(pdf_metadata),
                content=text_content,
                metadata={
                    **pdf_metadata,
                    "file_type": "pdf",
                    "page_count": len(reader.pages),
                    "char_count": len(text_content),
                }
            )
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse PDF: {str(e)}") from e

    def supports_format(self, filename: str) -> bool:
        """
        Check if parser supports the file format.

        Args:
            filename: Name of the file to check

        Returns:
            bool: True if PDF format is supported
        """
        return filename.lower().endswith('.pdf')

    def extract_metadata(self, raw: BinaryIO) -> dict[str, Any]:
        """
        Extract metadata from raw PDF bytes.

        Args:
            raw: Binary PDF content

        Returns:
            dict: Extracted metadata
        """
        if not pypdf:
            return {}
        
        try:
            reader = pypdf.PdfReader(raw)
            return self._extract_pdf_metadata(reader, None)
        except Exception:
            return {}

    def _extract_pdf_metadata(self, reader: Any, user_metadata: dict[str, Any] | None) -> dict[str, Any]:
        """
        Extract metadata from PDF reader.

        Args:
            reader: PyPDF2 reader object
            user_metadata: User-provided metadata

        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata: dict[str, Any] = {}
        
        # Basic PDF info
        if reader.pages:
            metadata.update({
                "page_count": len(reader.pages),
                "is_encrypted": getattr(reader, 'is_encrypted', False),
                "has_xfa": getattr(reader, '_has_xfa', False),
                "has_xfa_signatures": getattr(reader, '_has_xfa_signatures', False),
            })
        
        # Document info
        if hasattr(reader, 'documentInfo') and reader.documentInfo:
            doc_info = reader.documentInfo
            if doc_info:
                metadata.update({
                    "title": doc_info.get('/Title'),
                    "author": doc_info.get('/Author'),
                    "subject": doc_info.get('/Subject'),
                    "creator": doc_info.get('/Creator'),
                    "producer": doc_info.get('/Producer'),
                    "creation_date": doc_info.get('/CreationDate'),
                    "modification_date": doc_info.get('/ModDate'),
                    "keywords": doc_info.get('/Keywords'),
                })
        
        # Include user metadata
        if user_metadata:
            metadata.update(user_metadata)
        
        return metadata

    def _extract_title_from_metadata(self, metadata: dict[str, Any]) -> str | None:
        """Extract title from PDF metadata."""
        # Try multiple fields for title
        title = metadata.get("title")
        if title:
            return str(title)
        
        title = metadata.get("subject")
        if title:
            return str(title)
        
        return None

    def _generate_document_id(self, metadata: dict[str, Any] | None) -> str:
        """Generate document ID from metadata or use default."""
        if metadata and "document_id" in metadata:
            return str(metadata["document_id"])
        
        import uuid
        return str(uuid.uuid4())


class TextDocumentParser(DocumentParser):
    """Plain text document parser."""

    def __init__(self):
        """Initialize text parser."""
        pass

    def parse(self, raw: BinaryIO, metadata: dict[str, Any] | None = None) -> Document:
        """
        Parse text document bytes into structured Document.

        Args:
            raw: Binary text content
            metadata: Optional metadata to attach

        Returns:
            Document: Structured document with text and metadata

        Raises:
            DocumentParsingError: If text parsing fails
        """
        try:
            # Seek to beginning
            raw.seek(0)
            
            # Read and decode text
            content = raw.read().decode('utf-8')
            
            # Extract metadata
            text_metadata = self._extract_text_metadata(content, metadata)
            
            return Document(
                id=self._generate_document_id(metadata),
                workspace_id=metadata.get("workspace_id", "default") if metadata else "default",
                title=metadata.get("title", self._extract_title_from_content(content)) if metadata else self._extract_title_from_content(content),
                content=content,
                metadata={
                    **text_metadata,
                    "file_type": "text",
                    "char_count": len(content),
                    "line_count": len(content.splitlines()),
                    "word_count": len(content.split()),
                }
            )
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse text: {str(e)}") from e

    def supports_format(self, filename: str) -> bool:
        """
        Check if parser supports the file format.

        Args:
            filename: Name of the file to check

        Returns:
            bool: True if text format is supported
        """
        return filename.lower().endswith('.txt')

    def extract_metadata(self, raw: BinaryIO) -> dict[str, Any]:
        """
        Extract metadata from raw text bytes.

        Args:
            raw: Binary text content

        Returns:
            dict: Extracted metadata
        """
        try:
            raw.seek(0)
            content = raw.read().decode('utf-8')
            return self._extract_text_metadata(content, None)
        except Exception:
            return {}

    def _extract_text_metadata(self, content: str, user_metadata: dict[str, Any] | None) -> dict[str, Any]:
        """
        Extract metadata from text content.

        Args:
            content: Text content
            user_metadata: User-provided metadata

        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata: dict[str, Any] = {
            "encoding": "utf-8",
        }
        
        # Basic text statistics
        if content:
            lines = content.splitlines()
            words = content.split()
            
            metadata.update({
                "line_count": len(lines),
                "word_count": len(words),
                "char_count": len(content),
                "has_content": len(content.strip()) > 0,
            })
            
            # Try to extract potential title from first line
            if lines and len(lines[0].strip()) > 0:
                first_line = lines[0].strip()
                if len(first_line) < 100:  # Reasonable title length
                    metadata["potential_title"] = first_line
        
        # Include user metadata
        if user_metadata:
            metadata.update(user_metadata)
        
        return metadata

    def _extract_title_from_content(self, content: str) -> str | None:
        """Extract title from text content."""
        lines = content.splitlines()
        
        # Look for first non-empty line as potential title
        for line in lines:
            stripped = line.strip()
            if stripped and len(stripped) < 100:  # Reasonable title length
                return stripped
        
        return None

    def _generate_document_id(self, metadata: dict[str, Any] | None) -> str:
        """Generate document ID from metadata or use default."""
        if metadata and "document_id" in metadata:
            return str(metadata["document_id"])
        
        import uuid
        return str(uuid.uuid4())


class DocxDocumentParser(DocumentParser):
    """DOCX document parser using python-docx library."""

    def __init__(self):
        """Initialize DOCX parser."""
        pass

    def parse(self, raw: BinaryIO, metadata: dict[str, Any] | None = None) -> Document:
        """
        Parse DOCX document bytes into structured Document.

        Args:
            raw: Binary DOCX content
            metadata: Optional metadata to attach

        Returns:
            Document: Structured document with text and metadata

        Raises:
            DocumentParsingError: If DOCX parsing fails
        """
        if not DocxDocument:
            raise DocumentParsingError("python-docx library not available for DOCX parsing")

        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(raw.read())
                temp_file_path = temp_file.name
            
            try:
                # Open the document
                doc = DocxDocument(temp_file_path)
                
                # Extract text from paragraphs
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                text_content = "\n".join(text_parts)
                
                # Extract metadata from DOCX
                docx_metadata = self._extract_docx_metadata(doc, metadata)
                
                return Document(
                    id=self._generate_document_id(metadata),
                    workspace_id=metadata.get("workspace_id", "default") if metadata else "default",
                    title=metadata.get("title", self._extract_title_from_metadata(docx_metadata)) if metadata else self._extract_title_from_metadata(docx_metadata),
                    content=text_content,
                    metadata={
                        **docx_metadata,
                        "file_type": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "char_count": len(text_content),
                    }
                )
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}") from e

    def supports_format(self, filename: str) -> bool:
        """
        Check if parser supports the file format.

        Args:
            filename: Name of the file to check

        Returns:
            bool: True if DOCX format is supported
        """
        return filename.lower().endswith('.docx')

    def extract_metadata(self, raw: BinaryIO) -> dict[str, Any]:
        """
        Extract metadata from raw DOCX bytes.

        Args:
            raw: Binary DOCX content

        Returns:
            dict: Extracted metadata
        """
        if not DocxDocument:
            return {}
        
        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(raw.read())
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                return self._extract_docx_metadata(doc, None)
            finally:
                os.unlink(temp_file_path)
        except Exception:
            return {}

    def _extract_docx_metadata(self, doc: Any, user_metadata: dict[str, Any] | None) -> dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: python-docx document object
            user_metadata: User-provided metadata

        Returns:
            Dict[str, Any]: Extracted metadata
        """
        metadata: dict[str, Any] = {}
        
        # Basic document info
        metadata.update({
            "paragraph_count": len(doc.paragraphs),
        })
        
        # Core properties
        if hasattr(doc, 'core_properties'):
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "category": core_props.category,
                    "comments": core_props.comments,
                    "created": core_props.created,
                    "modified": core_props.modified,
                    "last_modified_by": core_props.last_modified_by,
                    "revision": core_props.revision,
                })
        
        # Include user metadata
        if user_metadata:
            metadata.update(user_metadata)
        
        return metadata

    def _extract_title_from_metadata(self, metadata: dict[str, Any]) -> str | None:
        """Extract title from DOCX metadata."""
        # Try multiple fields for title
        title = metadata.get("title")
        if title:
            return str(title)
        
        title = metadata.get("subject")
        if title:
            return str(title)
        
        return None

    def _generate_document_id(self, metadata: dict[str, Any] | None) -> str:
        """Generate document ID from metadata or use default."""
        if metadata and "document_id" in metadata:
            return str(metadata["document_id"])
        
        import uuid
        return str(uuid.uuid4())


class DocumentParsingError(Exception):
    """Exception raised when document parsing fails."""
    pass