"""DOCX document parser implementation."""

import os
import tempfile
import uuid
from typing import TYPE_CHECKING, BinaryIO, Optional

from returns.result import Failure, Result, Success

from src.infrastructure.rag.steps.general.parsing.document_parser import (
    DocumentParser,
    ParsingError,
)
from src.infrastructure.types.common import MetadataDict
from src.infrastructure.types.document import Document

if TYPE_CHECKING:
    from docx import Document as DocxDocumentType

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False


class DocxDocumentParser(DocumentParser):
    """DOCX document parser using python-docx library."""

    def parse(
        self, raw: BinaryIO, metadata: Optional[MetadataDict] = None
    ) -> Result[Document, ParsingError]:
        """
        Parse DOCX document bytes into structured Document.

        Args:
            raw: Binary DOCX content
            metadata: Optional metadata to attach

        Returns:
            Result containing Document on success, or ParsingError on failure
        """
        if not DOCX_AVAILABLE or DocxDocument is None:
            return Failure(
                ParsingError(
                    "python-docx library not available. Install: pip install python-docx",
                    code="DEPENDENCY_ERROR",
                )
            )

        temp_file_path: Optional[str] = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(raw.read())
                temp_file_path = temp_file.name

            doc = DocxDocument(temp_file_path)

            text_parts: list[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            text_content = "\n".join(text_parts)

            docx_metadata = self._extract_docx_metadata(doc, metadata)
            doc_id = self._generate_document_id(metadata)
            workspace_id = str(metadata.get("workspace_id", "default")) if metadata else "default"
            title = self._get_title(metadata, docx_metadata) or "Untitled Document"
            return Success(
                Document(
                    id=doc_id,
                    workspace_id=workspace_id,
                    title=title,
                    content=text_content,
                    metadata={
                        "file_type": "docx",
                        "paragraph_count": str(len(doc.paragraphs)),
                        "char_count": str(len(text_content)),
                    },
                )
            )

        except Exception as e:
            return Failure(ParsingError(f"Failed to parse DOCX: {e}", code="PARSE_ERROR"))
        finally:
            if temp_file_path:
                try:
                    os.unlink(temp_file_path)
                except OSError:
                    pass

    def supports_format(self, filename: str) -> bool:
        """Check if parser supports the file format."""
        return filename.lower().endswith(".docx")

    def extract_metadata(self, raw: BinaryIO) -> MetadataDict:
        """Extract metadata from raw DOCX bytes."""
        if not DOCX_AVAILABLE or DocxDocument is None:
            return {}

        temp_file_path: Optional[str] = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(raw.read())
                temp_file_path = temp_file.name

            doc = DocxDocument(temp_file_path)
            return self._extract_docx_metadata(doc, None)
        except Exception:
            return {}
        finally:
            if temp_file_path:
                try:
                    os.unlink(temp_file_path)
                except OSError:
                    pass

    def _extract_docx_metadata(
        self, doc: "DocxDocumentType", user_metadata: Optional[MetadataDict]
    ) -> MetadataDict:
        """Extract metadata from DOCX document."""
        metadata: MetadataDict = {
            "paragraph_count": len(doc.paragraphs),
        }

        if hasattr(doc, "core_properties"):
            core_props_data = self._extract_core_properties(doc.core_properties)
            metadata.update(core_props_data)

        if user_metadata:
            metadata.update(user_metadata)

        return metadata

    def _extract_core_properties(
        self, core_props: "docx_module.opc.coreprops.CoreProperties"
    ) -> MetadataDict:
        """Extract core properties and return metadata."""
        result: MetadataDict = {}

        if not core_props:
            return result

        if core_props.title:
            result["title"] = str(core_props.title)
        if core_props.author:
            result["author"] = str(core_props.author)
        if core_props.subject:
            result["subject"] = str(core_props.subject)
        if core_props.keywords:
            result["keywords"] = str(core_props.keywords)
        if core_props.category:
            result["category"] = str(core_props.category)

        return result

    def _get_title(self, metadata: Optional[MetadataDict], docx_metadata: MetadataDict) -> Optional[str]:
        """Get title from metadata or DOCX metadata."""
        if metadata and "title" in metadata:
            return str(metadata["title"])

        title = docx_metadata.get("title")
        if title:
            return str(title)

        subject = docx_metadata.get("subject")
        if subject:
            return str(subject)

        return None

    def _generate_document_id(self, metadata: Optional[MetadataDict]) -> str:
        """Generate document ID from metadata or use default."""
        if metadata and "document_id" in metadata:
            return str(metadata["document_id"])
        return str(uuid.uuid4())
